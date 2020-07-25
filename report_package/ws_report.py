import new_logger as lg
from datetime import datetime, time, date, timedelta
import os
import sys
import pandas as pd
from db_package import db_ops
from docx import Document
from docx.shared import Inches

class ws_rp():
    def __init__(self):
        wx = lg.get_handle()
        wx.info("[OBJ] ws_rp __init__ called")
        try:
            self.db = db_ops(host='127.0.0.1', db='stock', user='wx', pwd='5171013')
        except Exception as e:
            raise e

    def __del__(self):
        wx = lg.get_handle()
        self.db.cursor.close()
        self.db.handle.close()
        wx.info("[OBJ] ws_rp : __del__() called")

    def output_docx(self, filename = 'null', para_dict = None):
        wx = lg.get_handle()
        work_path = os.path.dirname(os.path.abspath(sys.argv[0]))
        output_path = work_path + '\\report\\'
        today = date.today().strftime('%Y%m%d')
        filename = output_path + today + "_" + filename + ".docx"

        # 创建文档对象
        doc = Document()

        # 设置文档标题，中文要用unicode字符串
        doc.add_heading(para_dict['title'], 0)
        para_dict.pop('title')

        for key in para_dict.keys():
            # 添加一级标题
            doc.add_heading(key, level=1)

            if para_dict[key].empty:
                p = doc.add_paragraph('没有数据 ')
                p.add_run('！ ').bold = True
                continue
            df = para_dict[key]
            # 添加表格
            table = doc.add_table(rows=df.shape[0]+1, cols=df.shape[1])
            hdr_cells = table.rows[0].cells
            col_counter = 0
            for col_name in df.columns.tolist():
                hdr_cells[col_counter].text = col_name
                col_counter += 1

            row_counter = 1
            for index, row in df.iterrows():
                col_counter = 0
                hdr_cells = table.rows[row_counter].cells
                for col in list(df.columns):
                    hdr_cells[col_counter].text = str(row[col])
                    col_counter += 1
                row_counter += 1

        doc.save(filename)
        wx.info("[output_docx] {} is outputed ".format(filename))

    def output_table(self, dd_df=None, filename='null', sheet_name=None, type='.xlsx', index=False):
        wx = lg.get_handle()
        work_path = os.path.dirname(os.path.abspath(sys.argv[0]))
        output_path = work_path + '\\report\\'
        today = date.today().strftime('%Y%m%d')
        filename = output_path + today + "_" + filename + type
        if dd_df is None or dd_df.empty:
            wx.info("[output_file] Dataframe is empty")
        elif sheet_name is None:
            sheet_name = 'Noname'
        else:
            if type == '.xlsx':
                dd_df.to_excel(filename,index=index, sheet_name= sheet_name, float_format="%.2f", encoding='utf_8_sig')
            else:
                dd_df.to_csv(filename, index=index, encoding='utf_8_sig')
        wx.info("[output_cvs] {} is completed".format(filename))

    def calc_total_amount(self):
        # wx = lg.get_handle()
        # start_date = (date.today() + timedelta(days=0)).strftime('%Y%m%d')
        # wx.info("get list a total amount called")
        # self.db.db_call_procedure("list_a_total_amount", start_date, 1, 2, 3, 4, 5, 6)
        # self.db.cursor.callproc(p_name, (args[0], args[1], args[2], args[3], args[4], args[5], args[6]))
        p_name = "list_a_total_amount"
        self.db.cursor.callproc(p_name, ('20190124', 1, 2, 3, 4, 5, 6, 7, 8))
        self.db.cursor.execute(
            "select @_" + p_name + "_0, @_" + p_name + "_1, @_" + p_name + "_2, @_" + p_name + "_3, @_"
            + p_name + "_4, @_" + p_name + "_5, @_" + p_name + "_6, @_" + p_name + "_7, @_" + p_name + "_8")
        result = self.db.cursor.fetchall()
        self.db.handle.commit()


    """
    # 大宗交易数据统计 ：股票代码、 交易次数、 成交量、 均价、最高价、最低价
    # 对比指定日期的 收盘价， 默认是当前日期的前一天收盘价
    """
    def ws_price_compare_close_price(self, days=180, close_date=None):
        wx = lg.get_handle()
        table_name = { '00': 'stock.code_00_201901', '60': 'stock.code_60_201901',
                       '30': 'stock.code_30_201901', '002': 'stock.code_002_201901'}

        start_date = (date.today() + timedelta(days=-days)).strftime('%Y%m%d')
        df_price_comparision = pd.DataFrame()

        if close_date is None:
            close_price_date = (date.today() + timedelta(days=-1)).strftime('%Y%m%d')
        else:
            close_price_date = close_date
        for index, type in enumerate(table_name):
            sql= "SELECT ws.id as `证券代码`, la.name as `名称`, count(ws.id) as `大宗交易次数`," \
                 " sum(ws.vol) as `大宗成交量（万股）` , dd.close as `收盘价("+close_price_date+")`, " \
                 " sum(ws.amount)/sum(ws.vol) as `大宗买入均价`, max(ws.price) as `大宗最高买价`,  min(ws.price) as `大宗最低买价` " \
                 " FROM ws_201901 as ws  " \
                 " left join "+table_name[type]+ "  as dd on dd.id = ws.id " \
                 " left join list_a as la on la.id=ws.id where ws.id like '"+type+\
                 "%' and ws.date > "+start_date+" and dd.date = "+close_price_date+\
                 "  group by ws.id order by `大宗交易次数` desc"
            iCount = self.db.cursor.execute(sql)
            self.db.handle.commit()
            if iCount > 0:
                # wx.info("[calc days vol] acquire {} rows of result".format(iCount))
                arr_record = self.db.cursor.fetchall()
                columnDes = self.db.cursor.description  # 获取连接对象的描述信息
                columnNames = [columnDes[i][0] for i in range(len(columnDes))]
                sub_df_price = pd.DataFrame([list(i) for i in arr_record], columns=columnNames)
                if df_price_comparision.empty:
                    df_price_comparision = sub_df_price
                else:
                    df_price_comparision = df_price_comparision.append(sub_df_price)
            else:
                wx.info("[ws_price_compare_close_price] failed to exec SQL {}".format(sql))
        return df_price_comparision

    def calc_days_vol(self, days=0, limit=10):
        wx = lg.get_handle()
        table_name = {'60': 'stock.code_60_201901', '30': 'stock.code_30_201901',
                      '002': 'stock.code_002_201901', '00': 'stock.code_00_201901'}
        df_days_vol = pd.DataFrame()

        for index, type in enumerate(table_name):
            # wx.info("{} , {}".format(index, type))
            sql = "SELECT dd.id as `证券代码`, la.list_date as 上市日期, la.name as 名称, sw1.industry_name as 申万一级行业, " \
                  "sw2.industry_name as 申万二级行业, sum(dd.vol) / 100 as `累计成交量(万股)`, " \
                  "la.flow_shares as `流动股本(万股)`, sum(dd.vol) / (flow_shares * 100) as `占比(%)`, " \
                  "sum(dd.vol) / (" + str(days) + "*100) as `日均成交量(万股)` " \
                  "FROM " + table_name[type] + " as dd  " \
                  "left join list_a as la on dd.id = la.id " \
                  "left join sw_industry_code as sw2 on la.sw_level_2 = sw2.industry_code " \
                  "left join sw_industry_code as sw1 on la.sw_level_1 = sw1.industry_code " \
                  "where str_to_date(dd.date, '%Y%m%d') > (current_date() - " + str(days) + ") " \
                  "group by dd.id order by `占比(%)` desc limit " + str(limit)

            iCount = self.db.cursor.execute(sql)
            self.db.handle.commit()
            if iCount > 0:
                # wx.info("[calc days vol] acquire {} rows of result".format(iCount))
                arr_days_vol = self.db.cursor.fetchall()
                columnDes = self.db.cursor.description  # 获取连接对象的描述信息
                columnNames = [columnDes[i][0] for i in range(len(columnDes))]
                sub_df_days_vol = pd.DataFrame([list(i) for i in arr_days_vol], columns=columnNames)
                if df_days_vol.empty:
                    df_days_vol = sub_df_days_vol
                else:
                    df_days_vol = df_days_vol.append(sub_df_days_vol)
            else:
                wx.info("[calc days vol] failed to exec SQL {}".format(sql))
        df_days_vol.sort_values(by="占比(%)", ascending=False, inplace=True)
        return df_days_vol



    def dgj_trading_summary(self, days=90):
        wx = lg.get_handle()
        start_date = (date.today() + timedelta(days=-days)).strftime('%Y%m%d')
        buy_sql = "SELECT dgj.id as `证券代码`, la.name as `名称`, count(dgj.id) as `高管买入次数`, " \
                  " sum(dgj.vol) /10000 as `高管买入（万股）`, sum(dgj.amount)/sum(dgj.vol) as `高管买入均价`, " \
                  " max(dgj.price) as `高管最高买价`, min(dgj.price) as `高管最低买价` " \
                  " FROM dgj_201901 as dgj left join list_a as la on dgj.id = la.id " \
                  " where dgj.date > "+start_date+" and dgj.vol > 0  group by dgj.id"

        sell_sql = "SELECT dgj.id as `证券代码`, la.name as `名称`, count(dgj.id) as `高管卖出次数`, " \
                   " sum(dgj.vol) /10000 as `高管卖出（万股）`, sum(dgj.amount)/sum(dgj.vol) as `高管卖出均价`," \
                   " max(dgj.price) as `高管最高卖价`, min(dgj.price) as `高管最低卖价` " \
                   " FROM dgj_201901 as dgj left join list_a as la on dgj.id = la.id " \
                   " where dgj.date > " + start_date + " and dgj.vol < 0  group by dgj.id"

        buy_df = self.db._exec_sql(buy_sql)
        sell_df = self.db._exec_sql(sell_sql)
        df_dgj_trading = pd.merge(buy_df, sell_df, how='outer', left_on=['证券代码','名称'], right_on=['证券代码','名称'])
        # df_dgj_trading.fillna(0, inplace=True)
        return  df_dgj_trading

"""
# SELECT id, sum(vol) as `成交量` , sum(amount)/sum(vol) as `机构买入均价`, max(price) as `最高价`, min(price) as `最低价`, count(id) as `机构交易次数` 
# FROM stock.ws_201901 where date > 20180501
# group by id order by `机构交易次数` desc ;
"""

# df = pd.read_excel(os.getcwd() + os.sep + 'stock.xlsx',converters = {u'证券代码':str})